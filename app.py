from flask import Flask, render_template, request, redirect, url_for, flash, send_file, abort
from flask_login import login_user, login_required, logout_user, current_user
from datetime import datetime
import os
from dotenv import load_dotenv
import requests
import json
import time
from extensions import db, login_manager, bcrypt
from models import User, Essay
from forms import LoginForm, RegistrationForm, ProfileForm
from werkzeug.utils import secure_filename
import io
import markupsafe
import markdown2
import docx  # Add this import for docx processing

# Load environment variables
load_dotenv()

# Allowed file extensions
ALLOWED_EXTENSIONS = {'txt', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def read_docx(file):
    doc = docx.Document(file)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():  # Only add non-empty paragraphs
            full_text.append(para.text.strip())
    return '\n'.join(full_text)

def read_txt(file):
    return file.read().decode('utf-8')

def create_app():
    app = Flask(__name__)
    app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'your-secret-key-here')
    app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL')
    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
    app.config['UPLOAD_FOLDER'] = os.path.join('static', 'uploads')
    app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

    # Initialize extensions
    db.init_app(app)
    bcrypt.init_app(app)
    login_manager.init_app(app)

    # Add context processor for datetime
    @app.context_processor
    def inject_now():
        return {'now': datetime.now()}

    # Add markdown filter
    @app.template_filter('markdown')
    def markdown_filter(text):
        if not text:
            return ""
        # Convert markdown to HTML with extras for both headings and breaks
        html = markdown2.markdown(text, extras=['break-on-newline', 'header-ids', 'tables', 'fenced-code-blocks'])
        # Clean up any excessive line breaks
        while '<br><br>' in html:
            html = html.replace('<br><br>', '<br>')
        return markupsafe.Markup(html)

    # Add clean_html_breaks filter
    @app.template_filter('clean_html_breaks')
    def clean_html_breaks_filter(s):
        if not s:
            return ""
        # Replace all variations of <br> tags with a single newline
        s = s.replace('<br><br>', '\n')
        s = s.replace('<br /><br />', '\n')
        s = s.replace('<br/><br/>', '\n')
        s = s.replace('<br>', '\n')
        s = s.replace('<br />', '\n')
        s = s.replace('<br/>', '\n')
        # Remove any duplicate newlines
        while '\n\n' in s:
            s = s.replace('\n\n', '\n')
        return s

    # Add nl2br filter
    @app.template_filter('nl2br')
    def nl2br_filter(s):
        if not s:
            return ""
        return markupsafe.Markup(s.replace('\n', '<br>'))

    # Add extract_score filter
    @app.template_filter('extract_score')
    def extract_score_filter(review_text, score_type):
        if not review_text:
            return "0"
        try:
            import re
            pattern = f"{score_type}:\\s*(\\d+(?:\\.\\d+)?)/100"
            match = re.search(pattern, review_text, re.IGNORECASE)
            if match:
                score = float(match.group(1))
                return f"{score:.1f}" if score % 1 != 0 else f"{int(score)}"
            return "0"
        except Exception as e:
            print(f"Error extracting score: {str(e)}")
            return "0"

    # Add extract_section filter
    @app.template_filter('extract_section')
    def extract_section_filter(review_text, section_name):
        if not review_text:
            return "No content available."
        try:
            import re
            pattern = f"{section_name}:(.*?)(?=(?:AI-Generated Score:|Humanized Score:|Total Score:|Detailed Feedback:|$))"
            match = re.search(pattern, review_text, re.IGNORECASE | re.DOTALL)
            if match:
                content = match.group(1).strip()
                return content if content else "No content available."
            return "No content available."
        except Exception as e:
            print(f"Error extracting section: {str(e)}")
            return "No content available."

    # Create upload folder if it doesn't exist
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

    # Create database tables
    with app.app_context():
        db.create_all()  # Only create tables if they don't exist

    # Configure OpenRouter API
    OPENROUTER_API_KEY = os.getenv('OPENROUTER_API_KEY')
    OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
    # List of OpenRouter free models to try in order
    MODEL_LIST = [
        "meta-llama/llama-4-maverick:free",
        "meta-llama/llama-4-scout:free",
        "nvidia/llama-3.1-nemotron-ultra-253b-v1:free",
        "bytedance-research/ui-tars-72b:free",
        "qwen/qwen3-32b:free",
        "deepseek/deepseek-chat-v3-0324:free",
        "moonshotai/kimi-vl-a3b-thinking:free",
        "mistralai/mistral-small-3.1-24b-instruct:free",
        "qwen/qwen2.5-vl-72b-instruct:free",
        "deepseek/deepseek-r1-distill-llama-70b:free"
    ]
    DEFAULT_MODEL = MODEL_LIST[0]

    # Rate limit handling
    MAX_RETRIES = 3
    RETRY_DELAY = 5  # seconds

    def handle_rate_limit_error(error_msg):
        """Return user-friendly message for rate limit errors"""
        return "API rate limit reached. Please try again in a few seconds. If this persists, you may need to wait a bit longer."

    def generate_with_retry(prompt, model=None, max_retries=MAX_RETRIES):
        """Try each model in MODEL_LIST until one works, with retry logic for rate limits"""
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost:5000",
            "X-Title": "ThinkDraft"
        }
        data = {
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        }
        last_error = None
        for model_name in MODEL_LIST:
            data["model"] = model_name
            for attempt in range(max_retries):
                try:
                    print(f"Attempting API call with model: {model_name}")
                    response = requests.post(
                        url=OPENROUTER_API_URL,
                        headers=headers,
                        data=json.dumps(data)
                    )
                    print(f"Response Status: {response.status_code}")
                    print(f"Response Headers: {response.headers}")
                    print(f"Response Body: {response.text[:500]}")
                    response.raise_for_status()
                    result = response.json()
                    if 'choices' in result and len(result['choices']) > 0:
                        return result['choices'][0]['message']['content']
                    else:
                        raise Exception("No response generated from the model")
                except requests.exceptions.HTTPError as e:
                    if e.response.status_code == 429 or 'rate-limited' in response.text or 'Provider returned error' in response.text:
                        last_error = f"Model {model_name} is rate-limited or unavailable. Trying next model."
                        print(last_error)
                        time.sleep(RETRY_DELAY)
                        continue
                    last_error = f"API Error: {str(e)}"
                    print(last_error)
                    break
                except Exception as e:
                    last_error = f"Error: {str(e)}"
                    print(last_error)
                    break
        raise Exception(last_error or "All models failed. Please try again later.")

    @login_manager.user_loader
    def load_user(user_id):
        return db.session.get(User, int(user_id))

    # Routes
    @app.route('/')
    def home():
        return render_template('home.html')

    @app.route('/register', methods=['GET', 'POST'])
    def register():
        if current_user.is_authenticated:
            flash('You are already logged in!', 'info')
            return redirect(url_for('dashboard'))
        form = RegistrationForm()
        if form.validate_on_submit():
            hashed_password = bcrypt.generate_password_hash(form.password.data).decode('utf-8')
            user = User(username=form.username.data, email=form.email.data, password=hashed_password)
            db.session.add(user)
            db.session.commit()
            flash('Your account has been created! You can now log in.', 'success')
            return redirect(url_for('login'))
        return render_template('register.html', title='Register', form=form)

    @app.route('/login', methods=['GET', 'POST'])
    def login():
        if current_user.is_authenticated:
            flash('You are already logged in!', 'info')
            return redirect(url_for('dashboard'))
        form = LoginForm()
        if form.validate_on_submit():
            user = User.query.filter_by(email=form.email.data).first()
            if user and bcrypt.check_password_hash(user.password, form.password.data):
                login_user(user, remember=form.remember.data)
                next_page = request.args.get('next')
                flash('Welcome back!', 'success')
                return redirect(next_page) if next_page else redirect(url_for('dashboard'))
            else:
                flash('Login unsuccessful. Please check email and password.', 'danger')
        return render_template('login.html', title='Login', form=form)

    @app.route('/logout')
    @login_required
    def logout():
        logout_user()
        flash('You have been logged out successfully.', 'info')
        return redirect(url_for('home'))

    @app.route('/dashboard')
    @login_required
    def dashboard():
        essays = Essay.query.filter_by(user_id=current_user.id).order_by(Essay.date_posted.desc()).all()
        return render_template('dashboard.html', essays=essays)

    @app.route('/generate', methods=['GET', 'POST'])
    @login_required
    def generate_essay():
        if request.method == 'POST':
            topic = request.form.get('topic')
            word_count = request.form.get('word_count')
            style = request.form.get('style', 'academic')
            additional_instructions = request.form.get('additional_instructions', '')
            
            try:
                prompt = f"""Write a {style} essay about {topic} with approximately {word_count} words.
                {additional_instructions}
                The essay should be well-structured, engaging, and professional.
                Include an introduction, body paragraphs with supporting evidence, and a conclusion.
                Use double line breaks between paragraphs for proper formatting."""
                
                content = generate_with_retry(prompt)
                
                # Clean up the content by replacing HTML tags with proper newlines
                content = (content
                    .replace('<br><br>', '\n\n')
                    .replace('<br /><br />', '\n\n')
                    .replace('<br/><br/>', '\n\n')
                    .replace('<br>', '\n')
                    .replace('<br />', '\n')
                    .replace('<br/>', '\n')
                    .strip())  # Remove any leading/trailing whitespace
                
                # Ensure consistent paragraph spacing
                paragraphs = [p.strip() for p in content.split('\n\n')]
                content = '\n\n'.join(p for p in paragraphs if p)
                
                new_essay = Essay(
                    title=topic,
                    content=content,
                    user_id=current_user.id
                )
                
                db.session.add(new_essay)
                db.session.commit()
                
                flash('Essay generated successfully!', 'success')
                return redirect(url_for('dashboard'))
            except Exception as e:
                print(f"Error details: {str(e)}")  # Print detailed error for debugging
                flash(str(e), 'danger')
                return redirect(url_for('generate_essay'))
        
        return render_template('generate.html')

    @app.route('/review/<int:essay_id>', methods=['GET', 'POST'])
    @login_required
    def review_essay(essay_id):
        essay = Essay.query.get_or_404(essay_id)
        
        if request.method == 'POST':
            try:
                prompt = f"""
Review the following essay and provide only numerical scores in the following format. Do not add any other text or explanations:

AI-Generated Score: XX/100
Humanized Score: XX/100
Total Score: XX/100
Quality Score: XX/100

Essay:
{essay.content}
"""
                # Get AI response
                response = generate_with_retry(prompt)
                print("Raw AI Response:", response)  # Debug print
                
                # Parse the response into structured format
                def parse_ai_response(response_text):
                    import re
                    
                    # Initialize structure
                    parsed = {
                        "ai_score": "0",
                        "human_score": "0",
                        "total_score": "0",
                        "quality_score": "0"
                    }
                    
                    try:
                        # Extract scores first
                        ai_score_match = re.search(r"AI-Generated Score:\s*(\d+(?:\.\d+)?)/100", response_text)
                        human_score_match = re.search(r"Humanized Score:\s*(\d+(?:\.\d+)?)/100", response_text)
                        total_score_match = re.search(r"Total Score:\s*(\d+(?:\.\d+)?)/100", response_text)
                        quality_score_match = re.search(r"Quality Score:\s*(\d+(?:\.\d+)?)/100", response_text)
                        
                        print("Score matches:", {  # Debug print
                            "ai": ai_score_match.group(1) if ai_score_match else "not found",
                            "human": human_score_match.group(1) if human_score_match else "not found",
                            "total": total_score_match.group(1) if total_score_match else "not found",
                            "quality": quality_score_match.group(1) if quality_score_match else "not found"
                        })
                        
                        if ai_score_match:
                            parsed["ai_score"] = ai_score_match.group(1)
                        if human_score_match:
                            parsed["human_score"] = human_score_match.group(1)
                        if total_score_match:
                            parsed["total_score"] = total_score_match.group(1)
                        if quality_score_match:
                            parsed["quality_score"] = quality_score_match.group(1)
                        
                        print("Parsed structure:", parsed)  # Debug print
                        
                    except Exception as e:
                        print(f"Error in parsing: {str(e)}")
                        raise
                    
                    # Calculate total score if not provided
                    if parsed["total_score"] == "0" and parsed["ai_score"] != "0" and parsed["human_score"] != "0":
                        ai_score = float(parsed["ai_score"])
                        human_score = float(parsed["human_score"])
                        total_score = (ai_score + human_score) / 2
                        parsed["total_score"] = f"{total_score:.1f}"
                    
                    # Format the structured review with proper line breaks
                    structured_review = f"""AI-Generated Score: {parsed['ai_score']}/100\nHumanized Score: {parsed['human_score']}/100\nTotal Score: {parsed['total_score']}/100\nQuality Score: {parsed['quality_score']}/100"""
                    
                    # Clean up any HTML tags and ensure consistent line breaks
                    structured_review = (structured_review
                        .replace('<br><br>', '\n\n')
                        .replace('<br /><br />', '\n\n')
                        .replace('<br/><br/>', '\n\n')
                        .replace('<br>', '\n')
                        .replace('<br />', '\n')
                        .replace('<br/>', '\n')
                        .strip())
                    
                    # Ensure consistent paragraph spacing
                    paragraphs = [p.strip() for p in structured_review.split('\n\n')]
                    structured_review = '\n\n'.join(p for p in paragraphs if p)
                    
                    return structured_review
                
                # Parse and structure the response
                structured_review = parse_ai_response(response)
                print("Final structured review:", structured_review)  # Debug print
                
                # Save the structured review
                essay.review = structured_review
                db.session.commit()
                
                flash('Essay reviewed successfully!', 'success')
                return redirect(url_for('dashboard'))
                
            except Exception as e:
                print(f"Error in review process: {str(e)}")  # Debug print
                flash(f"Error generating review: {str(e)}", 'danger')
                return redirect(url_for('review_essay', essay_id=essay_id))
        
        return render_template('review.html', essay=essay)

    @app.route('/profile', methods=['GET', 'POST'])
    @login_required
    def profile():
        form = ProfileForm()
        if form.validate_on_submit():
            if form.username.data != current_user.username:
                user = User.query.filter_by(username=form.username.data).first()
                if user:
                    flash('That username is already taken.', 'danger')
                    return redirect(url_for('profile'))
            
            current_user.username = form.username.data
            current_user.bio = form.bio.data
            
            if form.profile_picture.data:
                try:
                    filename = secure_filename(f"user_{current_user.id}_{form.profile_picture.data.filename}")
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    form.profile_picture.data.save(filepath)
                    current_user.profile_picture = filename
                except Exception as e:
                    flash(f'Error uploading profile picture: {str(e)}', 'danger')
                    return redirect(url_for('profile'))
            
            try:
                db.session.commit()
                flash('Your profile has been updated!', 'success')
            except Exception as e:
                db.session.rollback()
                flash(f'Error updating profile: {str(e)}', 'danger')
            
            return redirect(url_for('profile'))
            
        elif request.method == 'GET':
            form.username.data = current_user.username
            form.bio.data = current_user.bio
        
        return render_template('profile.html', title='Profile', form=form)

    @app.route('/delete_essay/<int:essay_id>', methods=['POST'])
    @login_required
    def delete_essay(essay_id):
        essay = Essay.query.get_or_404(essay_id)
        if essay.user_id != current_user.id:
            flash('You do not have permission to delete this essay.', 'danger')
            return redirect(url_for('dashboard'))
        db.session.delete(essay)
        db.session.commit()
        flash('Essay has been deleted.', 'success')
        return redirect(url_for('dashboard'))

    @app.route('/download_essay/<int:essay_id>/<format>')
    @login_required
    def download_essay(essay_id, format):
        essay = Essay.query.get_or_404(essay_id)
        if essay.user_id != current_user.id:
            flash('You do not have permission to download this essay.', 'danger')
            return redirect(url_for('dashboard'))

        try:
            if format == 'txt':
                return send_file(
                    io.BytesIO(essay.content.encode()),
                    mimetype='text/plain',
                    as_attachment=True,
                    download_name=f"{essay.title}.txt"
                )
            elif format == 'docx':
                try:
                    from docx import Document
                    doc = Document()
                    doc.add_heading(essay.title, 0)
                    doc.add_paragraph(essay.content)
                    
                    docx_file = io.BytesIO()
                    doc.save(docx_file)
                    docx_file.seek(0)
                    
                    return send_file(
                        docx_file,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        as_attachment=True,
                        download_name=f"{essay.title}.docx"
                    )
                except ImportError:
                    flash('Please install python-docx package to download as DOCX.', 'warning')
                    return redirect(url_for('dashboard'))
            else:
                flash('Invalid format specified.', 'danger')
                return redirect(url_for('dashboard'))
        except Exception as e:
            flash(f'Error downloading essay: {str(e)}', 'danger')
            return redirect(url_for('dashboard'))

    @app.route('/preview/<int:essay_id>')
    @login_required
    def preview_essay(essay_id):
        essay = Essay.query.get_or_404(essay_id)
        if essay.user_id != current_user.id:
            flash('You do not have permission to view this essay.', 'danger')
            return redirect(url_for('dashboard'))
        return render_template('preview.html', essay=essay)

    @app.route('/upload_review', methods=['GET', 'POST'])
    @login_required
    def upload_review():
        if request.method == 'POST':
            # Check if a file was uploaded
            if 'file' not in request.files:
                flash('No file selected', 'danger')
                return redirect(request.url)
            
            file = request.files['file']
            
            # If user submits without selecting a file
            if file.filename == '':
                flash('No file selected', 'danger')
                return redirect(request.url)
            
            if file and allowed_file(file.filename):
                try:
                    # Read the content based on file type
                    if file.filename.endswith('.docx'):
                        content = read_docx(file)
                    else:  # txt file
                        content = read_txt(file)
                    
                    # Create a new essay
                    title = os.path.splitext(secure_filename(file.filename))[0]
                    new_essay = Essay(
                        title=title,
                        content=content,
                        user_id=current_user.id
                    )
                    
                    db.session.add(new_essay)
                    db.session.commit()
                    
                    # Generate review immediately
                    try:
                        prompt = f"""
Review the following essay and provide scores and feedback in the following format:

AI-Generated Score: XX/100
Humanized Score: XX/100
Total Score: XX/100
Quality Score: XX/100

Essay:
{content}
"""
                        # Get AI response
                        response = generate_with_retry(prompt)
                        
                        # Save the review
                        new_essay.review = response
                        db.session.commit()
                        
                        flash('File uploaded and reviewed successfully!', 'success')
                        return redirect(url_for('preview_essay', essay_id=new_essay.id))
                    
                    except Exception as e:
                        flash(f'Error generating review: {str(e)}', 'danger')
                        return redirect(url_for('preview_essay', essay_id=new_essay.id))
                
                except Exception as e:
                    flash(f'Error processing file: {str(e)}', 'danger')
                    return redirect(request.url)
            else:
                flash('Invalid file type. Only .txt and .docx files are allowed.', 'danger')
                return redirect(request.url)
        
        return render_template('upload_review.html')

    @app.route('/privacy-policy')
    def privacy_policy():
        return render_template('privacy_policy.html', current_date=datetime.now().strftime('%B %d, %Y'))

    @app.route('/terms-of-service')
    def terms_of_service():
        return render_template('terms_of_service.html', current_date=datetime.now().strftime('%B %d, %Y'))

    @app.route('/view-review/<int:essay_id>')
    @login_required
    def view_review(essay_id):
        essay = Essay.query.get_or_404(essay_id)
        if essay.user_id != current_user.id:
            abort(403)
        if not essay.review:
            flash('This essay has not been reviewed yet.', 'warning')
            return redirect(url_for('dashboard'))
        return render_template('view_review.html', essay=essay)

    return app

DEBUG_MODE = os.getenv("DEBUG_NODE") == 'True'

if __name__ == '__main__':
    app = create_app()
    app.run(debug=DEBUG_MODE) 
